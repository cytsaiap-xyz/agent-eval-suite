#!/usr/bin/env python3
"""
Generate the sample Word document for P11.
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import random


def create_sample_document():
    doc = Document()

    # Set document properties
    doc.core_properties.author = "Jane Smith"
    doc.core_properties.title = "Q4 2024 Business Performance Report"
    doc.core_properties.created = datetime(2024, 1, 10)
    doc.core_properties.modified = datetime(2024, 1, 15)

    # Title
    title = doc.add_heading("Q4 2024 Business Performance Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report provides a comprehensive analysis of the company's performance "
        "during the fourth quarter of 2024. Key highlights include a 15% increase in "
        "revenue, successful expansion into three new markets, and the launch of our "
        "flagship product line. Despite challenging market conditions, the company "
        "maintained strong operational efficiency and customer satisfaction metrics."
    )

    # Introduction with inconsistent heading (intentional issue)
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Background", level=2)
    doc.add_paragraph(
        "The fourth quarter of 2024 marked a significant period for our organization. "
        "Building on the momentum from previous quarters, we executed our strategic "
        "initiatives while adapting to evolving market dynamics."
    )

    # Inconsistent heading style (intentional issue)
    p = doc.add_paragraph("1.2 Objectives")
    p.style = doc.styles['Heading 3']  # Wrong level - should be Heading 2

    doc.add_paragraph("The primary objectives for Q4 2024 were:")
    doc.add_paragraph("Achieve 12% revenue growth", style='List Bullet')
    doc.add_paragraph("Expand market presence in Asia-Pacific", style='List Bullet')
    doc.add_paragraph("Launch Product X before year-end", style='List Bullet')
    doc.add_paragraph("Maintain customer satisfaction above 90%", style='List Bullet')

    # Methodology
    doc.add_heading("2. Methodology", level=1)
    doc.add_paragraph(
        "This analysis employs a mixed-methods approach combining quantitative "
        "financial data with qualitative market research. Data was collected from "
        "internal systems, customer surveys, and third-party market reports."
    )

    # Add a table
    doc.add_heading("2.1 Data Sources", level=2)
    table1 = doc.add_table(rows=5, cols=3)
    table1.style = 'Table Grid'

    # Header row
    headers = ["Source", "Type", "Coverage"]
    for i, header in enumerate(headers):
        table1.rows[0].cells[i].text = header

    # Data rows
    data = [
        ["ERP System", "Financial", "All departments"],
        ["CRM Platform", "Customer", "Sales & Support"],
        ["Market Reports", "External", "Industry trends"],
        ["Employee Surveys", "Internal", "All staff"],
    ]
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table1.rows[i].cells[j].text = cell_data

    # Results section
    doc.add_heading("3. Results", level=1)

    doc.add_heading("3.1 Financial Performance", level=2)
    doc.add_paragraph(
        "The company achieved strong financial results in Q4 2024, exceeding "
        "initial projections across most key metrics."
    )

    # Financial results table
    doc.add_heading("3.1.1 Revenue Analysis", level=3)
    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'

    # Header
    fin_headers = ["Metric", "Q4 2023", "Q4 2024", "Change"]
    for i, h in enumerate(fin_headers):
        table2.rows[0].cells[i].text = h

    fin_data = [
        ["Revenue", "$12.5M", "$14.4M", "+15.2%"],
        ["Gross Margin", "42%", "44%", "+2pp"],
        ["Operating Income", "$2.1M", "$2.8M", "+33.3%"],
        ["Net Income", "$1.5M", "$2.0M", "+33.3%"],
        ["EPS", "$0.75", "$1.00", "+33.3%"],
    ]
    for i, row_data in enumerate(fin_data, 1):
        for j, cell_data in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_data

    doc.add_paragraph("")  # Spacer
    doc.add_paragraph(
        "Note: All financial figures are unaudited and subject to final review."
    ).italic = True

    doc.add_heading("3.2 Market Expansion", level=2)
    doc.add_paragraph(
        "We successfully entered three new markets during Q4:"
    )
    doc.add_paragraph("Singapore - Launched November 2024", style='List Number')
    doc.add_paragraph("Vietnam - Launched December 2024", style='List Number')
    doc.add_paragraph("Thailand - Soft launch December 2024", style='List Number')

    doc.add_heading("3.3 Product Performance", level=2)
    doc.add_paragraph(
        "Product X was launched on schedule in November 2024, achieving "
        "significant early traction with over 10,000 units sold in the first month."
    )

    # Missing heading level (intentional issue) - jumps from 2 to 4
    doc.add_heading("3.3.1.1 Product X Sales Breakdown", level=4)

    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Table Grid'
    prod_headers = ["Region", "Units Sold", "Revenue"]
    for i, h in enumerate(prod_headers):
        table3.rows[0].cells[i].text = h

    prod_data = [
        ["North America", "6,500", "$975K"],
        ["Europe", "2,800", "$420K"],
        ["Asia-Pacific", "700", "$105K"],
    ]
    for i, row_data in enumerate(prod_data, 1):
        for j, cell_data in enumerate(row_data):
            table3.rows[i].cells[j].text = cell_data

    # Customer section
    doc.add_heading("4. Customer Analysis", level=1)

    doc.add_heading("4.1 Satisfaction Metrics", level=2)
    doc.add_paragraph(
        "Customer satisfaction remained strong throughout Q4, with our Net "
        "Promoter Score (NPS) reaching an all-time high of 72."
    )

    doc.add_heading("4.2 Support Performance", level=2)
    doc.add_paragraph("Key support metrics:")
    doc.add_paragraph("Average response time: 2.3 hours (target: <4 hours)", style='List Bullet')
    doc.add_paragraph("First contact resolution: 78% (target: >75%)", style='List Bullet')
    doc.add_paragraph("Customer effort score: 4.2/5 (target: >4.0)", style='List Bullet')

    # Challenges
    doc.add_heading("5. Challenges and Risks", level=1)
    doc.add_paragraph(
        "Despite strong overall performance, several challenges emerged during Q4:"
    )
    doc.add_paragraph("Supply chain disruptions affecting component availability", style='List Bullet')
    doc.add_paragraph("Increased competition in core markets", style='List Bullet')
    doc.add_paragraph("Currency fluctuations impacting international revenue", style='List Bullet')
    doc.add_paragraph("Talent acquisition in key technical roles", style='List Bullet')

    # Recommendations
    doc.add_heading("6. Recommendations", level=1)
    doc.add_paragraph("Based on Q4 performance and market analysis, we recommend:")
    doc.add_paragraph("Accelerate investment in Asia-Pacific region", style='List Number')
    doc.add_paragraph("Expand Product X to enterprise segment", style='List Number')
    doc.add_paragraph("Implement additional supply chain redundancy", style='List Number')
    doc.add_paragraph("Increase R&D budget by 20% for FY2025", style='List Number')

    # Conclusion
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "Q4 2024 demonstrated the company's ability to execute effectively while "
        "navigating a complex market environment. The strong financial results, "
        "successful market expansion, and positive customer metrics position the "
        "company well for continued growth in 2025. Key focus areas for the coming "
        "year include scaling operations in new markets, expanding the product "
        "portfolio, and maintaining operational excellence."
    )

    # Appendix
    doc.add_heading("Appendix A: Detailed Financial Statements", level=1)
    doc.add_paragraph(
        "Detailed financial statements are available upon request from the "
        "Finance department."
    )

    doc.add_heading("Appendix B: Glossary", level=1)
    doc.add_paragraph("EPS: Earnings Per Share", style='List Bullet')
    doc.add_paragraph("NPS: Net Promoter Score", style='List Bullet')
    doc.add_paragraph("pp: Percentage points", style='List Bullet')
    doc.add_paragraph("FY: Fiscal Year", style='List Bullet')

    # Save
    doc.save("report_draft.docx")
    print("Created report_draft.docx")


if __name__ == "__main__":
    create_sample_document()
